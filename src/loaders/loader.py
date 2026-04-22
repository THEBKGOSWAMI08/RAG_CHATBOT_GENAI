import os
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_core.documents import Document


def load_pdf(file_path: str) -> list[Document]:
    loader = PyPDFLoader(file_path)
    return loader.load()


def load_txt(file_path: str) -> list[Document]:
    loader = TextLoader(file_path, encoding="utf-8")
    return loader.load()


def load_docx(file_path: str) -> list[Document]:
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    return [Document(page_content=text, metadata={"source": file_path})]


LOADER_MAP = {
    ".pdf": load_pdf,
    ".txt": load_txt,
    ".docx": load_docx,
}


def load_documents(path: str) -> list[Document]:
    """Load documents from a file or directory."""
    documents = []

    if os.path.isfile(path):
        ext = os.path.splitext(path)[1].lower()
        loader_fn = LOADER_MAP.get(ext)
        if loader_fn:
            documents.extend(loader_fn(path))
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    elif os.path.isdir(path):
        for root, _, files in os.walk(path):
            for file in files:
                ext = os.path.splitext(file)[1].lower()
                loader_fn = LOADER_MAP.get(ext)
                if loader_fn:
                    file_path = os.path.join(root, file)
                    documents.extend(loader_fn(file_path))
    else:
        raise FileNotFoundError(f"Path not found: {path}")

    return documents
#code  editied 